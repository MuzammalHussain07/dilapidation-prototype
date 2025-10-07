# app.py
import streamlit as st
from backend import analyze_image_bytes
from docx import Document
from docx.shared import Inches
import io

# 🌟 App title
st.title("Dilapidation Auto Detection Prototype")

# 📤 Upload multiple images
uploaded = st.file_uploader(
    "Upload one or more images", 
    type=["jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

# Process uploaded images
if uploaded:
    st.info("Processing images, please wait...")
    results = []
    
    for f in uploaded:
        # Read image bytes
        img_bytes = f.read()
        
        # Analyze the image using your trained model
        out_pil, summary = analyze_image_bytes(io.BytesIO(img_bytes))
        
        # Show image + result
        st.image(out_pil, caption=f"{f.name} — {summary.get('description')}", use_column_width=True)
        st.write(summary)
        
        results.append((f.name, out_pil, summary))

    # 📝 Optional: Generate Word report
    if st.button("📄 Generate Word Report"):
        doc = Document()
        doc.add_heading("Dilapidation Survey Report", 0)

        for i, (name, pil_img, summary) in enumerate(results, start=1):
            doc.add_heading(f"Issue {i}: {name}", level=2)
            doc.add_paragraph(summary.get("description", ""))
            
            # Convert image to bytes for Word file
            img_io = io.BytesIO()
            pil_img.save(img_io, format="JPEG")
            img_io.seek(0)
            doc.add_picture(img_io, width=Inches(5))
        
        # Save and provide download link
        out_file = "dilapidation_report.docx"
        doc.save(out_file)
        
        with open(out_file, "rb") as fh:
            st.download_button("📥 Download Word Report", fh, file_name=out_file)

else:
    st.info("👆 Please upload images above to begin analysis.")
